"""
文档图片提取测试脚本

本脚本用于测试从Word文档中提取图片的功能，并验证提取顺序是否正确。
特别针对"AI接口测试系统建设方案.docx"文件进行测试。
"""

import os
import sys
import uuid
from pathlib import Path
from typing import List, Dict, Any

# 导入所需的库
from docling.document_converter import DocumentConverter
from docling.datamodel.document import DoclingDocument
from docling_core.types.doc import PictureItem

# 导入python-docx库
try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.parts.image import ImagePart
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx库未安装，无法使用python-docx方法提取图片")

class ImageExtractionTester:
    """测试从Word文档中提取图片的功能"""
    
    def __init__(self, temp_dir: str = "./test_extracted_images"):
        """初始化测试器"""
        self.temp_image_dir = temp_dir
        os.makedirs(self.temp_image_dir, exist_ok=True)
        self.doc_converter = DocumentConverter()
    
    def extract_images_from_document(self, document: DoclingDocument, original_file_path: str = None) -> List[Dict[str, Any]]:
        """
        从文档中提取所有图片
        
        Args:
            document: 文档对象
            original_file_path: 原始文件路径
            
        Returns:
            图片信息列表，每个元素包含图片路径和在文档中的位置信息
        """
        # 提取文档中的所有图片
        images_info = []
        picture_counter = 0
        extraction_method = "unknown"
        
        # 调试信息
        print("开始提取图片...")
        
        # 根据文档类型选择最佳提取方法
        if original_file_path:
            file_ext = os.path.splitext(original_file_path.lower())[1]
            
            # 对于Word文档，使用python-docx方法
            if file_ext == '.docx' and DOCX_AVAILABLE:
                print("使用python-docx库提取Word文档图片...")
                try:
                    extraction_method = "python-docx"
                    doc = docx.Document(original_file_path)
                    
                    # 方法1: 通过遍历文档内容来确定图片顺序
                    print("使用文档内容遍历方法确定图片顺序...")
                    
                    # 存储所有图片关系
                    image_rels = {}
                    for rel_id, rel in doc.part.rels.items():
                        if isinstance(rel.target_part, ImagePart):
                            image_rels[rel_id] = rel
                    
                    # 遍历文档中的所有段落和表格，查找图片引用
                    ordered_rel_ids = []
                    
                    # 检查文档主体中的图片
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            # 检查run的XML内容中的图片引用
                            xml = run._element.xml
                            for rel_id in image_rels.keys():
                                # 在XML中查找图片关系ID
                                if rel_id in xml and rel_id not in ordered_rel_ids:
                                    ordered_rel_ids.append(rel_id)
                    
                    # 检查表格中的图片
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        xml = run._element.xml
                                        for rel_id in image_rels.keys():
                                            if rel_id in xml and rel_id not in ordered_rel_ids:
                                                ordered_rel_ids.append(rel_id)
                    
                    # 添加任何可能在遍历过程中遗漏的图片关系
                    for rel_id in image_rels.keys():
                        if rel_id not in ordered_rel_ids:
                            ordered_rel_ids.append(rel_id)
                    
                    print(f"找到的图片关系ID顺序: {ordered_rel_ids}")
                    
                    # 按照找到的顺序提取图片
                    for doc_idx, rel_id in enumerate(ordered_rel_ids):
                        rel = image_rels[rel_id]
                        # 获取图片数据
                        image_data = rel.target_part.blob
                        
                        # 保存图片到临时文件
                        picture_counter += 1
                        image_filename = f"docx_image_{doc_idx}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(image_data)
                        
                        # 记录图片信息
                        images_info.append({
                            "image_path": image_path,
                            "rel_id": rel_id,
                            "document_index": doc_idx,
                            "position": picture_counter,
                            "extraction_method": extraction_method
                        })
                except Exception as e:
                    print(f"使用python-docx提取图片时出错: {str(e)}")
                    print(f"错误详情: {type(e).__name__}")
                    import traceback
                    traceback.print_exc()
            else:
                print(f"不支持的文件类型: {file_ext}，或python-docx库未安装")
        else:
            print("未提供原始文件路径，无法提取图片")
        
        print(f"总共提取到 {len(images_info)} 张图片，使用方法: {extraction_method}")
        return images_info
    
    def test_image_extraction(self, file_path: str):
        """
        测试从文档中提取图片
        
        Args:
            file_path: 文档路径
        """
        print(f"测试文件: {file_path}")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"错误: 文件 {file_path} 不存在")
            return
        
        # 转换文档
        conversion_result = self.doc_converter.convert(file_path)
        document = conversion_result.document
        
        # 提取图片
        images_info = self.extract_images_from_document(document, original_file_path=file_path)
        
        # 打印提取结果
        print(f"\n=== 提取结果 ===")
        print(f"从文件 {file_path} 中提取到 {len(images_info)} 张图片")
        
        # 打印每张图片的详细信息
        for i, image_info in enumerate(images_info):
            print(f"\n图片 {i+1}:")
            print(f"  路径: {image_info['image_path']}")
            print(f"  关系ID: {image_info['rel_id']}")
            print(f"  文档索引: {image_info['document_index']}")
            print(f"  位置: {image_info['position']}")
            print(f"  提取方法: {image_info['extraction_method']}")
        
        # 生成HTML报告，方便查看图片
        self._generate_html_report(file_path, images_info)
    
    def _generate_html_report(self, file_path: str, images_info: List[Dict[str, Any]]):
        """
        生成HTML报告，显示提取的图片及其顺序
        
        Args:
            file_path: 文档路径
            images_info: 图片信息列表
        """
        file_name = os.path.basename(file_path)
        report_path = os.path.join(self.temp_image_dir, "extraction_report.html")
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>图片提取报告 - {file_name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #333; }}
                .image-container {{ 
                    border: 1px solid #ddd; 
                    margin: 10px 0; 
                    padding: 10px; 
                    border-radius: 5px;
                }}
                .image-info {{ margin-bottom: 10px; }}
                img {{ max-width: 100%; max-height: 300px; border: 1px solid #eee; }}
            </style>
        </head>
        <body>
            <h1>图片提取报告 - {file_name}</h1>
            <p>总共提取到 {len(images_info)} 张图片</p>
        """
        
        for i, image_info in enumerate(images_info):
            image_path = image_info['image_path']
            rel_id = image_info['rel_id']
            doc_idx = image_info['document_index']
            position = image_info['position']
            
            # 获取相对路径，用于HTML中显示图片
            rel_image_path = os.path.relpath(image_path, self.temp_image_dir)
            
            html_content += f"""
            <div class="image-container">
                <div class="image-info">
                    <h2>图片 {i+1}</h2>
                    <p><strong>文档索引:</strong> {doc_idx}</p>
                    <p><strong>关系ID:</strong> {rel_id}</p>
                    <p><strong>位置:</strong> {position}</p>
                    <p><strong>文件名:</strong> {os.path.basename(image_path)}</p>
                </div>
                <img src="{rel_image_path}" alt="图片 {i+1}">
            </div>
            """
        
        html_content += """
        </body>
        </html>
        """
        
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        print(f"\n生成了HTML报告: {report_path}")
        print(f"请在浏览器中打开此文件，查看提取的图片及其顺序")


# 测试代码
if __name__ == "__main__":
    # 设置测试文件路径
    default_file_path = "AI接口测试系统建设方案.docx"
    
    # 如果命令行提供了文件路径，则使用命令行参数
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = default_file_path
    
    # 创建测试器并运行测试
    tester = ImageExtractionTester()
    tester.test_image_extraction(file_path)
