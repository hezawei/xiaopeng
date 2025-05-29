"""
文档图片处理器示例

本模块演示如何处理Word文档中的图片，包括：
1. 提取文档中的所有图片
2. 使用多模态大模型(Moonshot)对图片进行详细描述
3. 将图片描述替换到文档中的相应位置
4. 输出包含图片描述的纯文本结果或新的Word文档

使用方法:
```python
processor = DocumentImageProcessor()
# 处理文档并输出纯文本
text_with_image_descriptions = processor.process_document("path/to/document.docx")
print(text_with_image_descriptions)

# 或处理文档并输出新的Word文档
processor.process_document_to_word("path/to/document.docx", "path/to/output.docx")
```
"""

import asyncio
import aiohttp
import base64
import os
import time
import uuid
import requests
import docx
from typing import List, Dict, Any, Optional, Tuple, TYPE_CHECKING

# 使用TYPE_CHECKING条件导入，避免运行时导入但允许类型检查
if TYPE_CHECKING:
    from docling.datamodel.document import DoclingDocument
else:
    DoclingDocument = Any  # 运行时使用Any类型

# 确保导入所需的库
try:
    from docling.datamodel.document import DoclingDocument
    from docling_core.types.doc import TextItem, PictureItem
    from docling.document_converter import DocumentConverter
except ImportError:
    print("警告: docling库未安装，某些功能可能不可用")
    DoclingDocument = object  # 定义为object而不是Any
    TextItem = object
    PictureItem = object
    DocumentConverter = object

class DocumentImageProcessor:
    """
    处理文档中图片的类
    
    该类提供了一套完整的文档图片处理流程，包括：
    1. 从Word文档中提取文本和图片
    2. 使用多模态大模型对图片进行详细描述
    3. 将图片描述替换到文档中的相应位置
    4. 输出包含图片描述的纯文本结果或Markdown文本
    """
    
    def __init__(
        self,
        api_keys: List[str] = None,
        api_base: str = "https://api.moonshot.cn/v1",
        multimodal_model: str = "moonshot-v1-32k-vision-preview",
        max_concurrent_requests: int = 4,  # 默认设为4，与API key数量匹配
        temp_dir: str = "./temp_images",
        min_api_interval: int = 0  # 移除强制等待间隔，改为动态处理
    ):
        """
        初始化文档图片处理器
        
        Args:
            api_keys: API密钥列表，用于轮换调用
            api_base: API基础URL
            multimodal_model: 多模态模型名称
            max_concurrent_requests: 最大并发请求数
            temp_dir: 临时图片存储目录
            min_api_interval: 同一API密钥最小调用间隔(秒)
        """
        # 使用新的API密钥列表（来自不同账号）
        if api_keys is None:
            self.api_keys = [
                "sk-fCgBeG8ETu8MMCIaAmdOI4JpOKKagF7qXNIE4kIhu2q8zEU3",  # 替换为您的第一个API密钥
                "sk-n9Kb2lFGyV3LJ0GEwg3NBjy5ZrJk3qTlNJJt4Kfkro7T8Fct",  # 替换为您的第二个API密钥
                "sk-GwCmhwrrhicbkbishgUhYmVg6IrAJkg4mwWPNIMpdkQBe8tr",  # 替换为您的第三个API密钥
                "sk-3vO3Ku08wFZBRREAWGxPWKYtEcUE8pZ1ResjWOa0RvBLz2aM",  # 替换为您的第四个API密钥
            ]
        else:
            self.api_keys = api_keys if isinstance(api_keys, list) else [api_keys]
            
        self.api_base = api_base
        self.multimodal_model = multimodal_model
        self.max_concurrent_requests = min(max_concurrent_requests, len(self.api_keys))
        self.temp_image_dir = temp_dir
        self.min_api_interval = min_api_interval
        
        # 确保临时目录存在
        os.makedirs(self.temp_image_dir, exist_ok=True)
        
        # 初始化文档转换器
        try:
            self.doc_converter = DocumentConverter()
        except Exception as e:
            print(f"初始化DocumentConverter失败: {str(e)}")
            self.doc_converter = None
        
        # API密钥管理 - 使用锁和状态跟踪
        self.api_key_locks = {key: asyncio.Lock() for key in self.api_keys}
        self.api_key_last_used = {key: 0 for key in self.api_keys}
        self.api_key_in_use = {key: False for key in self.api_keys}  # 跟踪密钥是否正在使用
        
        # 创建API密钥池信号量 - 限制总并发数
        self.api_key_semaphore = asyncio.Semaphore(len(self.api_keys))
    
    def process_document(self, file_path: str) -> str:
        """
        处理文档，提取图片并生成描述，输出纯文本结果
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含图片描述的纯文本结果
        """
        print(f"开始处理文档: {file_path}")
        
        # 1. 转换文档并提取图片
        conversion_result = self.doc_converter.convert(file_path)
        document = conversion_result.document
        
        # 打印文档结构，帮助调试
        print(f"文档结构: {type(document)}")
        
        # 提取图片，传入原始文件路径
        images_info = self._extract_images_from_document(document, original_file_path=file_path)
        print(f"提取到 {len(images_info)} 张图片")
        
        if len(images_info) == 0:
            print("警告: 未从文档中提取到任何图片，请检查文档格式或docling库的兼容性")
            return document.export_to_markdown()
        
        # 2. 对每个图片生成描述
        for image_info in images_info:
            image_path = image_info["image_path"]
            print(f"正在描述图片: {image_path}")
            
            # 使用智能提示词让模型自行判断图片类型并给出相应描述
            description = self._describe_image_with_model(
                image_path,
                """请分析这张图片并提供详细描述。根据图片类型给出适当的描述:
                
                - 如果是流程图或架构图: 详细描述各个组件、流程步骤、数据流向和逻辑关系
                - 如果是表格: 描述表格结构、列名和主要数据内容
                - 如果是图表(如柱状图、饼图等): 解释图表表达的数据关系、趋势和关键数据点
                - 如果是界面截图: 描述界面布局、功能区域和主要控件
                - 如果是其他类型图片: 描述图片中的主要视觉元素和内容
                
                请确保描述全面、准确，覆盖图片中的所有重要信息，包括任何文本内容。"""
            )
            
            # 保存描述到图片信息中
            image_info["description"] = description
        
        # 3. 生成包含图片描述的文本
        result_text = self._generate_text_with_descriptions(document, images_info)
        
        print(f"文档处理完成，生成了包含{len(images_info)}个图片描述的文本")
        return result_text
    
    def process_document_to_word(self, file_path: str, output_path: str) -> str:
        """
        处理文档，提取图片并生成描述，输出新的Word文档
        
        Args:
            file_path: 输入文档路径
            output_path: 输出文档路径
            
        Returns:
            输出文档路径
        """
        print(f"开始处理文档: {file_path}")
        
        # 1. 转换文档并提取图片
        conversion_result = self.doc_converter.convert(file_path)
        document = conversion_result.document
        
        # 打印文档结构，帮助调试
        print(f"文档结构: {type(document)}")
        
        # 提取图片，传入原始文件路径
        images_info = self._extract_images_from_document(document, original_file_path=file_path)
        print(f"提取到 {len(images_info)} 张图片")
        
        if len(images_info) == 0:
            print("警告: 未从文档中提取到任何图片，请检查文档格式或docling库的兼容性")
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(document.export_to_markdown())
            return output_path
        
        # 2. 对每个图片生成描述
        for image_info in images_info:
            image_path = image_info["image_path"]
            print(f"正在描述图片: {image_path}")
            
            # 使用智能提示词让模型自行判断图片类型并给出相应描述
            description = self._describe_image_with_model(
                image_path,
                """请分析这张图片并提供详细描述。根据图片类型给出适当的描述:
                
                - 如果是流程图或架构图: 详细描述各个组件、流程步骤、数据流向和逻辑关系
                - 如果是表格: 描述表格结构、列名和主要数据内容
                - 如果是图表(如柱状图、饼图等): 解释图表表达的数据关系、趋势和关键数据点
                - 如果是界面截图: 描述界面布局、功能区域和主要控件
                - 如果是其他类型图片: 描述图片中的主要视觉元素和内容
                
                请确保描述全面、准确，覆盖图片中的所有重要信息，包括任何文本内容。"""
            )
            
            # 保存描述到图片信息中
            image_info["description"] = description
        
        # 3. 修改文档，将图片描述添加到文档中
        modified_document = self._modify_document_with_descriptions(document, images_info)
        
        # 4. 保存修改后的文档
        self._save_document_as_word(modified_document, output_path)
        
        print(f"文档处理完成，生成了包含{len(images_info)}个图片描述的Word文档: {output_path}")
        return output_path
    
    def _extract_images_from_document(self, document: "DoclingDocument", original_file_path: str = None) -> List[Dict[str, Any]]:
        """
        从文档中提取所有图片
        
        Args:
            document: 文档对象
            original_file_path: 原始文件路径
            
        Returns:
            图片信息列表，每个元素包含图片路径和在文档中的位置信息
        """
        # 提取文档中的所有图片
        images_info = []
        picture_counter = 0
        
        # 调试信息
        print("开始提取图片...")
        
        # 方法1: 尝试直接从DoclingDocument对象中提取图片
        if hasattr(document, 'pages'):
            print("尝试从document.pages提取图片...")
            for page_idx, page in enumerate(document.pages):
                for item in page.items:
                    if isinstance(item, PictureItem):
                        picture_counter += 1
                        # 保存图片到临时文件
                        image_filename = f"image_{page_idx}_{picture_counter}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        # 保存图片数据
                        with open(image_path, "wb") as f:
                            f.write(item.image_data)
                        
                        # 记录图片信息
                        images_info.append({
                            "image_path": image_path,
                            "page_idx": page_idx,
                            "item": item,
                            "position": picture_counter
                        })
    
        # 方法2: 尝试使用python-docx库直接从原始文档提取
        if len(images_info) == 0 and original_file_path and original_file_path.lower().endswith('.docx'):
            print("尝试使用python-docx库提取图片...")
            try:
                import docx
                from docx.document import Document as DocxDocument
                from docx.parts.image import ImagePart
                
                # 使用传入的原始文件路径
                doc = docx.Document(original_file_path)
                
                # 提取所有图片
                for rel_id, rel in doc.part.rels.items():
                    if isinstance(rel.target_part, ImagePart):
                        # 获取图片数据
                        image_data = rel.target_part.blob
                        
                        # 保存图片到临时文件
                        picture_counter += 1
                        image_filename = f"docx_image_{picture_counter}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(image_data)
                        
                        # 记录图片信息
                        images_info.append({
                            "image_path": image_path,
                            "rel_id": rel_id,
                            "position": picture_counter
                        })
            except Exception as e:
                print(f"使用python-docx提取图片时出错: {str(e)}")
    
        # 方法3: 尝试使用docling的原始转换结果
        if len(images_info) == 0 and hasattr(document, '_conversion_result'):
            print("尝试从conversion_result提取图片...")
            try:
                conversion_result = document._conversion_result
                if hasattr(conversion_result, 'images') and conversion_result.images:
                    for idx, img_data in enumerate(conversion_result.images):
                        picture_counter += 1
                        image_filename = f"conversion_image_{picture_counter}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(img_data)
                        
                        images_info.append({
                            "image_path": image_path,
                            "index": idx,
                            "position": picture_counter
                        })
            except Exception as e:
                print(f"从conversion_result提取图片时出错: {str(e)}")
    
        # 方法4: 如果文档是PDF，尝试使用PyMuPDF提取图片
        if len(images_info) == 0 and original_file_path and original_file_path.lower().endswith('.pdf'):
            print("尝试使用PyMuPDF提取图片...")
            try:
                import fitz  # PyMuPDF
                
                doc = fitz.open(original_file_path)
                
                for page_idx, page in enumerate(doc):
                    image_list = page.get_images(full=True)
                    
                    for img_idx, img_info in enumerate(image_list):
                        xref = img_info[0]
                        base_image = doc.extract_image(xref)
                        image_data = base_image["image"]
                        
                        picture_counter += 1
                        image_filename = f"pdf_image_{page_idx}_{img_idx}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(image_data)
                        
                        images_info.append({
                            "image_path": image_path,
                            "page_idx": page_idx,
                            "img_idx": img_idx,
                            "position": picture_counter
                        })
                
                doc.close()
            except Exception as e:
                print(f"使用PyMuPDF提取图片时出错: {str(e)}")
    
        # 方法5: 尝试直接从docx文件中提取图片（使用zipfile）
        if len(images_info) == 0 and original_file_path and original_file_path.lower().endswith('.docx'):
            print("尝试使用zipfile直接从docx文件中提取图片...")
            try:
                import zipfile
                from xml.etree import ElementTree as ET
                
                # 打开docx文件（实际上是一个zip文件）
                with zipfile.ZipFile(original_file_path) as docx_zip:
                    # 获取所有文件列表
                    file_list = docx_zip.namelist()
                    
                    # 查找所有图片文件（通常在word/media/目录下）
                    image_files = [f for f in file_list if f.startswith('word/media/')]
                    
                    for img_idx, img_file in enumerate(image_files):
                        # 提取图片数据
                        image_data = docx_zip.read(img_file)
                        
                        picture_counter += 1
                        image_filename = f"zip_image_{img_idx}_{uuid.uuid4().hex[:8]}.png"
                        image_path = os.path.join(self.temp_image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(image_data)
                        
                        images_info.append({
                            "image_path": image_path,
                            "source_file": img_file,
                            "position": picture_counter
                        })
            except Exception as e:
                print(f"使用zipfile提取图片时出错: {str(e)}")
    
        print(f"总共提取到 {len(images_info)} 张图片")
        return images_info



    def _describe_image_with_model(self, image_path: str, prompt: str) -> str:
        """
        使用多模态模型描述图片
        
        Args:
            image_path: 图片路径
            prompt: 描述提示词
            
        Returns:
            图片描述
        """
        try:
            # 读取图片并转换为base64
            with open(image_path, "rb") as image_file:
                image_data = base64.b64encode(image_file.read()).decode("utf-8")
            
            # 构建API请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            # 使用Moonshot多模态API
            payload = {
                "model": self.multimodal_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 2000
            }
            
            # 发送请求，使用循环而不是递归来处理重试
            max_retries = 5
            retry_count = 0
            retry_delay = 2  # 初始重试延迟（秒）
            
            while retry_count < max_retries:
                # 发送请求
                response = requests.post(
                    f"{self.api_base}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=60  # 设置较长的超时时间
                )
                
                # 解析响应
                if response.status_code == 200:
                    result = response.json()
                    description = result["choices"][0]["message"]["content"]
                    return description
                elif response.status_code == 429:
                    # 遇到速率限制，等待后重试
                    print(f"API请求受到速率限制，等待{retry_delay}秒后重试... ({retry_count+1}/{max_retries})")
                    time.sleep(retry_delay)
                    retry_delay *= 1.5  # 增加重试延迟
                    retry_count += 1
                else:
                    print(f"API请求失败: {response.status_code} {response.text}")
                    return f"[图片描述生成失败: API错误 {response.status_code}]"
            
            return "[图片描述生成失败: 超过最大重试次数]"
                
        except Exception as e:
            print(f"图片描述生成失败: {str(e)}")
            return f"[图片描述生成失败: {str(e)}]"
    
    def _generate_text_with_descriptions(self, document: DoclingDocument, images_info: List[Dict[str, Any]]) -> str:
        """
        生成包含图片描述的文本
        
        Args:
            document: 文档对象
            images_info: 图片信息列表
            
        Returns:
            包含图片描述的纯文本
        """
        # 获取文档的Markdown文本
        markdown_text = document.export_to_markdown()
        
        # 在Markdown文本中查找并替换图片标记
        result_text = markdown_text
        
        # 按照图片在文档中的位置排序
        sorted_images = sorted(images_info, key=lambda x: x["position"])
        
        # 替换图片标记
        for image_info in sorted_images:
            # 在Markdown中，图片通常表示为 ![alt text](image_url)
            # 在docling生成的Markdown中，可能是 <image> 或其他标记
            # 这里需要根据实际情况调整替换逻辑
            
            # 尝试多种可能的图片标记
            image_markers = ["<image>", "![", "<img"]
            replaced = False
            
            for marker in image_markers:
                if marker in result_text:
                    # 找到标记的位置
                    marker_pos = result_text.find(marker)
                    
                    # 如果是![或<img，需要找到标记的结束位置
                    if marker == "![":
                        # 查找下一个)，这是Markdown图片语法的结束
                        end_pos = result_text.find(")", marker_pos)
                        if end_pos > marker_pos:
                            # 替换整个图片标记
                            old_text = result_text[marker_pos:end_pos+1]
                            new_text = f"\n[图片描述: {image_info['description']}]\n"
                            result_text = result_text.replace(old_text, new_text, 1)
                            replaced = True
                            break
                    elif marker == "<img":
                        # 查找下一个>，这是HTML图片标签的结束
                        end_pos = result_text.find(">", marker_pos)
                        if end_pos > marker_pos:
                            # 替换整个图片标记
                            old_text = result_text[marker_pos:end_pos+1]
                            new_text = f"\n[图片描述: {image_info['description']}]\n"
                            result_text = result_text.replace(old_text, new_text, 1)
                            replaced = True
                            break
                    else:
                        # 直接替换标记
                        result_text = result_text.replace(
                            marker, 
                            f"\n[图片描述: {image_info['description']}]\n", 
                            1  # 只替换第一次出现的标记
                        )
                        replaced = True
                        break
            
            # 如果没有找到任何标记，则在文档末尾添加图片描述
            if not replaced:
                result_text += f"\n\n[图片 {image_info['position']} 描述: {image_info['description']}]\n"
        
        return result_text
    
    def _modify_document_with_descriptions(self, document: "DoclingDocument", images_info: List[Dict[str, Any]]) -> "DoclingDocument":
        """
        修改文档，将图片描述添加到文档中
        
        Args:
            document: 文档对象
            images_info: 图片信息列表
        
        Returns:
            修改后的文档对象
        """
        # 获取文档的Markdown文本
        markdown_text = document.export_to_markdown()
        
        # 创建一个新的文档内容，包含图片描述
        modified_text = markdown_text
        
        # 添加图片描述部分
        modified_text += "\n\n## 文档图片描述\n\n"
        
        # 为每个图片添加描述
        for i, image_info in enumerate(images_info):
            image_path = image_info["image_path"]
            description = image_info.get("description", "[无图片描述]")
            
            # 尝试确定图片在文档中的位置
            position_info = ""
            if "rel_id" in image_info:
                position_info = f"(文档中的第{i+1}张图片)"
            elif "page_idx" in image_info:
                position_info = f"(第{image_info['page_idx']+1}页的图片)"
            else:
                position_info = f"(图片{i+1})"
            
            # 添加图片描述
            modified_text += f"### 图片{i+1} {position_info}\n\n"
            modified_text += f"{description}\n\n"
            print(f"添加了图片描述: 图片{i+1} {position_info}")
        
        # 创建一个新的文档对象
        # 由于无法直接修改DoclingDocument对象，我们返回原始文档，但在保存时使用修改后的文本
        document._modified_text = modified_text  # 添加一个自定义属性存储修改后的文本
        
        return document
    
    def _find_parent_elements(self, document: DoclingDocument, element_id: str) -> Optional[Tuple[Any, int]]:
        """
        查找元素的父元素和索引
        
        Args:
            document: 文档对象
            element_id: 元素ID
            
        Returns:
            父元素和索引的元组，如果未找到则返回None
        """
        # 遍历文档中的所有元素
        for parent in document.iterate_containers():
            for i, child in enumerate(parent.children):
                if hasattr(child, 'id') and child.id == element_id:
                    return parent, i
        return None
    
    def _save_document_as_word(self, document: "DoclingDocument", output_path: str) -> None:
        """
        将文档保存为Word文档
        
        Args:
            document: 文档对象
            output_path: 输出文档路径
        """
        try:
            # 使用修改后的文本（如果存在）
            markdown_text = getattr(document, '_modified_text', document.export_to_markdown())
            
            # 获取文件扩展名
            ext = os.path.splitext(output_path)[1].lower()
            
            if ext == '.docx':
                try:
                    # 使用python-docx创建Word文档
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    
                    doc = Document()
                    
                    # 设置文档默认字体为宋体，解决中文显示问题
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = '宋体'
                    font.size = Pt(12)
                    
                    # 分割文本为段落和标题
                    paragraphs = markdown_text.split('\n\n')
                    
                    # 添加文本内容
                    for paragraph_text in paragraphs:
                        if not paragraph_text.strip():
                            continue
                        
                        # 处理Markdown标题
                        if paragraph_text.startswith('# '):
                            p = doc.add_heading(paragraph_text[2:], level=1)
                        elif paragraph_text.startswith('## '):
                            p = doc.add_heading(paragraph_text[3:], level=2)
                        elif paragraph_text.startswith('### '):
                            p = doc.add_heading(paragraph_text[4:], level=3)
                        else:
                            p = doc.add_paragraph(paragraph_text)
                    
                    # 保存文档
                    doc.save(output_path)
                    print(f"已将文档保存为Word文档: {output_path}")
                    return
                except Exception as e:
                    print(f"保存为Word文档失败，将保存为纯文本: {str(e)}")
            
            # 如果无法保存为Word或者输出路径不是.docx，则保存为纯文本
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(markdown_text)
            
            print(f"已将文档保存为纯文本: {output_path}")
        
        except Exception as e:
            print(f"保存文档失败: {str(e)}")
            
            # 尝试至少保存文本内容
            try:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write("文档保存失败，但以下是提取的文本内容:\n\n")
                    f.write(document.export_to_markdown())
                print(f"已将文本内容保存到: {output_path}")
            except Exception as inner_e:
                print(f"保存文本内容也失败了: {str(inner_e)}")

    async def process_document_async(self, file_path: str) -> str:
        """
        异步处理文档，提取图片并生成描述，输出纯文本结果
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含图片描述的纯文本结果
        """
        print(f"开始异步处理文档: {file_path}")
        
        # 1. 转换文档并提取图片
        conversion_result = self.doc_converter.convert(file_path)
        document = conversion_result.document
        
        # 提取图片，传入原始文件路径
        images_info = self._extract_images_from_document(document, original_file_path=file_path)
        print(f"提取到 {len(images_info)} 张图片")
        
        if len(images_info) == 0:
            print("警告: 未从文档中提取到任何图片，请检查文档格式或docling库的兼容性")
            return document.export_to_markdown()
        
        # 2. 创建所有图片描述任务
        prompt = """请分析这张图片并提供详细描述。根据图片类型给出适当的描述:
        
        - 如果是流程图或架构图: 详细描述各个组件、流程步骤、数据流向和逻辑关系
        - 如果是表格: 描述表格结构、列名和主要数据内容
        - 如果是图表(如柱状图、饼图等): 解释图表表达的数据关系、趋势和关键数据点
        - 如果是界面截图: 描述界面布局、功能区域和主要控件
        - 如果是其他类型图片: 描述图片中的主要视觉元素和内容
        
        请确保描述全面、准确，覆盖图片中的所有重要信息，包括任何文本内容。"""
        
        # 创建任务列表
        tasks = []
        for i, image_info in enumerate(images_info):
            image_path = image_info["image_path"]
            print(f"创建图片描述任务 {i+1}/{len(images_info)}: {image_path}")
            task = self._process_single_image(image_info, image_path, prompt)
            tasks.append(task)
        
        # 并发执行所有任务
        print(f"开始并发处理 {len(tasks)} 个图片描述任务")
        await asyncio.gather(*tasks)
        
        # 3. 生成包含图片描述的文本
        result_text = self._generate_text_with_descriptions(document, images_info)
        
        print(f"文档处理完成，生成了包含{len(images_info)}个图片描述的文本")
        return result_text

    async def _process_single_image(self, image_info: Dict[str, Any], image_path: str, prompt: str) -> None:
        """
        处理单个图片的描述任务
        
        Args:
            image_info: 图片信息字典，将被原地修改添加描述
            image_path: 图片路径
            prompt: 描述提示词
        """
        try:
            description = await self._describe_image_with_api_key_pool(image_path, prompt)
            image_info["description"] = description
            print(f"成功获取图片描述: {image_path}")
        except Exception as e:
            print(f"描述图片失败: {str(e)}")
            image_info["description"] = f"[图片描述生成失败: {str(e)}]"

    async def _describe_image_with_api_key_pool(self, image_path: str, prompt: str) -> str:
        """
        使用API密钥池异步描述图片
        
        Args:
            image_path: 图片路径
            prompt: 描述提示词
            
        Returns:
            图片描述
        """
        # 从API密钥池获取可用的API密钥
        async with self.api_key_semaphore:
            # 获取可用的API密钥
            api_key = await self._get_available_api_key()
            
            try:
                # 标记API密钥为使用中
                self.api_key_in_use[api_key] = True
                
                # 使用选中的API密钥发送请求
                return await self._send_image_description_request(image_path, prompt, api_key)
            finally:
                # 无论成功失败，都标记API密钥为可用
                self.api_key_in_use[api_key] = False
                # 更新最后使用时间
                self.api_key_last_used[api_key] = time.time()
    
    async def _get_available_api_key(self) -> str:
        """
        获取当前可用的API密钥，优先选择未被使用的密钥
        
        Returns:
            API密钥
        """
        # 首先尝试获取未被使用的密钥
        available_keys = [key for key, in_use in self.api_key_in_use.items() if not in_use]
        
        if available_keys:
            # 如果有未使用的密钥，按最后使用时间排序
            sorted_keys = sorted([(key, self.api_key_last_used[key]) for key in available_keys], 
                                key=lambda x: x[1])
            return sorted_keys[0][0]
        else:
            # 如果所有密钥都在使用中，等待任意一个密钥变为可用
            while True:
                await asyncio.sleep(0.1)  # 短暂等待
                available_keys = [key for key, in_use in self.api_key_in_use.items() if not in_use]
                if available_keys:
                    sorted_keys = sorted([(key, self.api_key_last_used[key]) for key in available_keys], 
                                        key=lambda x: x[1])
                    return sorted_keys[0][0]
    
    async def _send_image_description_request(self, image_path: str, prompt: str, api_key: str) -> str:
        """
        发送图片描述请求
        
        Args:
            image_path: 图片路径
            prompt: 描述提示词
            api_key: API密钥
            
        Returns:
            图片描述
        """
        print(f"使用API密钥 {api_key[:8]}... 处理图片: {image_path}")
        
        try:
            # 读取图片并转换为base64
            with open(image_path, "rb") as image_file:
                image_data = base64.b64encode(image_file.read()).decode("utf-8")
            
            # 构建API请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # 使用Moonshot多模态API
            payload = {
                "model": self.multimodal_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 2000
            }
            
            # 使用aiohttp发送异步请求
            async with aiohttp.ClientSession() as session:
                max_retries = 3
                retry_count = 0
                retry_delay = 2
                
                while retry_count < max_retries:
                    try:
                        async with session.post(
                            f"{self.api_base}/chat/completions",
                            headers=headers,
                            json=payload,
                            timeout=60
                        ) as response:
                            if response.status == 200:
                                result = await response.json()
                                description = result["choices"][0]["message"]["content"]
                                print(f"成功获取图片描述: {image_path}")
                                return description
                            elif response.status == 429:
                                # 速率限制，检查响应头中的Retry-After
                                retry_after = response.headers.get('Retry-After')
                                wait_time = int(retry_after) if retry_after and retry_after.isdigit() else retry_delay
                                
                                response_text = await response.text()
                                print(f"API密钥 {api_key[:8]}... 请求受到速率限制，等待{wait_time}秒后重试")
                                await asyncio.sleep(wait_time)
                                retry_delay = max(retry_delay * 1.5, wait_time * 1.2)  # 动态调整重试延迟
                                retry_count += 1
                            else:
                                response_text = await response.text()
                                raise Exception(f"API错误 {response.status}: {response_text}")
                    except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                        print(f"请求出错: {str(e)}，等待{retry_delay}秒后重试")
                        await asyncio.sleep(retry_delay)
                        retry_delay *= 1.5
                        retry_count += 1
                
                raise Exception("超过最大重试次数")
                
        except Exception as e:
            raise Exception(f"描述图片失败: {str(e)}")

    async def _describe_image_with_semaphore(self, semaphore, image_path: str, prompt: str) -> str:
        """
        使用信号量控制并发的图片描述方法
        
        Args:
            semaphore: 控制并发的信号量
            image_path: 图片路径
            prompt: 描述提示词
            
        Returns:
            图片描述
        """
        async with semaphore:
            print(f"开始处理图片: {image_path}")
            return await self._describe_image_with_model_async(image_path, prompt)

    async def process_document_to_text(self, file_path: str) -> str:
        """
        处理文档，提取图片并生成描述，输出纯文本结果
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含图片描述的纯文本结果
        """
        print(f"开始处理文档: {file_path}")
        
        # 1. 转换文档并提取图片
        conversion_result = self.doc_converter.convert(file_path)
        document = conversion_result.document
        
        # 打印文档结构，帮助调试
        print(f"文档结构: {type(document)}")
        
        # 提取图片，传入原始文件路径
        images_info = self._extract_images_from_document(document, original_file_path=file_path)
        print(f"提取到 {len(images_info)} 张图片")
        
        if len(images_info) == 0:
            print("警告: 未从文档中提取到任何图片，返回原始文档内容")
            return document.export_to_markdown()
        
        # 2. 创建所有图片描述任务
        prompt = """请分析这张图片并提供详细描述。根据图片类型给出适当的描述:
        
        - 如果是流程图或架构图: 详细描述各个组件、流程步骤、数据流向和逻辑关系
        - 如果是表格: 描述表格结构、列名和主要数据内容
        - 如果是图表(如柱状图、饼图等): 解释图表表达的数据关系、趋势和关键数据点
        - 如果是界面截图: 描述界面布局、功能区域和主要控件
        - 如果是其他类型图片: 描述图片中的主要视觉元素和内容
        
        请确保描述全面、准确，覆盖图片中的所有重要信息，包括任何文本内容。"""
        
        # 创建任务列表
        tasks = []
        for i, image_info in enumerate(images_info):
            image_path = image_info["image_path"]
            print(f"创建图片描述任务 {i+1}/{len(images_info)}: {image_path}")
            task = self._process_single_image(image_info, image_path, prompt)
            tasks.append(task)
        
        # 并发执行所有任务
        print(f"开始并发处理 {len(tasks)} 个图片描述任务")
        await asyncio.gather(*tasks)
        
        # 3. 生成包含图片描述的文本
        markdown_text = document.export_to_markdown()
        
        # 4. 将图片标记替换为图片描述
        result_text = self._replace_images_with_descriptions(markdown_text, images_info)
        
        print(f"文档处理完成，生成了包含{len(images_info)}个图片描述的文本")
        return result_text

    def _replace_images_with_descriptions(self, markdown_text: str, images_info: List[Dict[str, Any]]) -> str:
        """
        将Markdown文本中的图片标记替换为图片描述
        
        Args:
            markdown_text: Markdown格式的文本
            images_info: 图片信息列表
        
        Returns:
            替换后的文本
        """
        result_text = markdown_text
        
        # 按照图片在文档中的位置排序
        sorted_images = sorted(images_info, key=lambda x: x.get("position", 0))
        
        # 尝试查找并替换图片标记
        image_markers = ["![", "<img", "<image>", "\\includegraphics"]
        
        # 如果无法找到图片标记，则在文档末尾添加图片描述部分
        has_replaced = False
        
        # 尝试替换图片标记
        for marker in image_markers:
            marker_count = result_text.count(marker)
            if marker_count > 0 and marker_count >= len(sorted_images):
                has_replaced = True
                # 有足够的标记可以替换
                for i, image_info in enumerate(sorted_images):
                    description = image_info.get("description", "[无图片描述]")
                    # 查找第i+1个标记的位置
                    start_pos = -1
                    for _ in range(i+1):
                        start_pos = result_text.find(marker, start_pos + 1)
                    
                    if start_pos >= 0:
                        # 查找标记结束位置
                        if marker == "![":
                            end_pos = result_text.find(")", start_pos)
                            if end_pos < 0:
                                end_pos = result_text.find("]", start_pos)
                        elif marker == "<img":
                            end_pos = result_text.find(">", start_pos)
                        elif marker == "<image>":
                            end_pos = start_pos + len("<image>")
                        elif marker == "\\includegraphics":
                            end_pos = result_text.find("}", start_pos)
                        
                        if end_pos > start_pos:
                            # 替换图片标记为描述
                            image_desc = f"\n\n[图片描述 {i+1}]: {description}\n\n"
                            result_text = result_text[:start_pos] + image_desc + result_text[end_pos+1:]
        
        # 如果没有找到足够的图片标记，则在文档末尾添加图片描述部分
        if not has_replaced:
            result_text += "\n\n## 文档图片描述\n\n"
            for i, image_info in enumerate(sorted_images):
                description = image_info.get("description", "[无图片描述]")
                position_info = ""
                if "rel_id" in image_info:
                    position_info = f"(文档中的第{i+1}张图片)"
                elif "page_idx" in image_info:
                    position_info = f"(第{image_info['page_idx']+1}页的图片)"
                else:
                    position_info = f"(图片{i+1})"
                
                result_text += f"### 图片{i+1} {position_info}\n\n{description}\n\n"
        
        return result_text

    async def _process_images_serially(self, images_info: List[Dict[str, Any]]) -> None:
        """
        串行处理图片描述，避免API速率限制
        
        Args:
            images_info: 图片信息列表
        """
        # 智能提示词
        prompt = """请分析这张图片并提供详细描述。根据图片类型给出适当的描述:
        
        - 如果是流程图或架构图: 详细描述各个组件、流程步骤、数据流向和逻辑关系
        - 如果是表格: 描述表格结构、列名和主要数据内容
        - 如果是图表(如柱状图、饼图等): 解释图表表达的数据关系、趋势和关键数据点
        - 如果是界面截图: 描述界面布局、功能区域和主要控件
        - 如果是其他类型图片: 描述图片中的主要视觉元素和内容
        
        请确保描述全面、准确，覆盖图片中的所有重要信息，包括任何文本内容。"""
        
        # 轮流使用API密钥
        key_index = 0
        
        # 串行处理每张图片
        for i, image_info in enumerate(images_info):
            image_path = image_info["image_path"]
            api_key = self.api_keys[key_index]
            
            print(f"处理图片 {i+1}/{len(images_info)}: {image_path} (使用API密钥 {key_index})")
            
            # 尝试处理图片，包含重试逻辑
            max_retries = 5
            retry_count = 0
            base_wait_time = 5  # 基础等待时间（秒）
            
            while retry_count < max_retries:
                try:
                    # 处理图片
                    description = await self._describe_image_serially(image_path, prompt, api_key)
                    image_info["description"] = description
                    print(f"成功描述图片: {image_path}")
                    
                    # 成功后等待固定时间，避免触发速率限制
                    await asyncio.sleep(base_wait_time)
                    break
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"描述图片失败: {error_msg}")
                    
                    # 解析错误信息中的等待时间建议
                    wait_time = base_wait_time
                    if "please try again after" in error_msg:
                        try:
                            # 尝试从错误消息中提取建议等待时间
                            import re
                            time_match = re.search(r"after (\d+) seconds", error_msg)
                            if time_match:
                                suggested_wait = int(time_match.group(1))
                                # 添加额外缓冲时间
                                wait_time = suggested_wait + 2
                        except:
                            # 如果提取失败，使用默认等待时间
                            pass
                    
                    # 指数退避策略
                    retry_count += 1
                    if retry_count < max_retries:
                        actual_wait = wait_time * (2 ** (retry_count - 1))
                        print(f"等待 {actual_wait} 秒后重试... ({retry_count}/{max_retries})")
                        await asyncio.sleep(actual_wait)
                        
                        # 轮换到下一个API密钥
                        key_index = (key_index + 1) % len(self.api_keys)
                        api_key = self.api_keys[key_index]
                        print(f"切换到API密钥 {key_index}")
                    else:
                        # 超过最大重试次数
                        image_info["description"] = f"[图片描述生成失败: 超过最大重试次数]"
            
            # 即使成功也轮换API密钥，分散负载
            key_index = (key_index + 1) % len(self.api_keys)

    async def _describe_image_serially(self, image_path: str, prompt: str, api_key: str) -> str:
        """
        串行处理单张图片描述
        
        Args:
            image_path: 图片路径
            prompt: 描述提示词
            api_key: API密钥
        
        Returns:
            图片描述
        """
        try:
            # 读取图片并转换为base64
            with open(image_path, "rb") as image_file:
                image_data = base64.b64encode(image_file.read()).decode("utf-8")
            
            # 构建API请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # 使用Moonshot多模态API
            payload = {
                "model": self.multimodal_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_data}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 2000
            }
            
            # 使用aiohttp发送请求
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.api_base}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=60  # 设置较长的超时时间
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        return result["choices"][0]["message"]["content"]
                    else:
                        response_text = await response.text()
                        raise Exception(f"API错误 {response.status}: {response_text}")
        
        except Exception as e:
            raise Exception(f"描述图片失败: {str(e)}")

    def process_for_llamaindex(self, file_path: str) -> str:
        """
        处理文档并返回适合LlamaIndex处理的文本内容
        
        Args:
            file_path: 文档路径
        
        Returns:
            处理后的文本内容
        """
        # 使用同步方式运行异步函数
        return asyncio.run(self.process_document_to_text(file_path))

# 示例用法
if __name__ == "__main__":
    import asyncio
    
    async def main():
        # 新的API密钥列表（来自不同账号）
        api_keys = [
            "sk-fCgBeG8ETu8MMCIaAmdOI4JpOKKagF7qXNIE4kIhu2q8zEU3",  # 替换为您的第一个API密钥
            "sk-n9Kb2lFGyV3LJ0GEwg3NBjy5ZrJk3qTlNJJt4Kfkro7T8Fct",  # 替换为您的第二个API密钥
            "sk-GwCmhwrrhicbkbishgUhYmVg6IrAJkg4mwWPNIMpdkQBe8tr",  # 替换为您的第三个API密钥
            "sk-3vO3Ku08wFZBRREAWGxPWKYtEcUE8pZ1ResjWOa0RvBLz2aM",  # 替换为您的第四个API密钥
        ]
        
        # 创建处理器实例
        processor = DocumentImageProcessor(
            api_keys=api_keys,
            api_base="https://api.moonshot.cn/v1",
            multimodal_model="moonshot-v1-32k-vision-preview",
            max_concurrent_requests=4  # 设置为API密钥数量
        )
        
        # 处理文档并输出文本
        document_path = "AI接口测试系统建设方案.docx"  # 替换为实际的文档路径
        output_text_path = "example_with_image_descriptions.txt"
        
        print(f"开始处理文档: {document_path}")
        
        # 使用异步方法处理文档
        result_text = await processor.process_document_to_text(document_path)
        
        # 保存结果到文件
        try:
            with open(output_text_path, "w", encoding="utf-8") as f:
                f.write(result_text)
            print(f"\n处理结果已保存到: {output_text_path}")
        except Exception as e:
            print(f"保存结果失败: {str(e)}")
            print("\n处理结果预览:")
            print(result_text[:500] + "...")

    # 运行异步主函数
    asyncio.run(main())






